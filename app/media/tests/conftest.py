"""
Test fixtures for media app.

Provides fixtures for:
- Sample files (JPEG, PNG, PDF, etc.)
- Invalid files (executable, empty, oversized)
- Authenticated test clients
- Users with storage quotas
"""

from __future__ import annotations

import io
from typing import TYPE_CHECKING

import pytest
from django.core.files.uploadedfile import SimpleUploadedFile
from PIL import Image
from rest_framework.test import APIClient
from rest_framework_simplejwt.tokens import RefreshToken

from authentication.tests.factories import UserFactory

if TYPE_CHECKING:
    from authentication.models import User


# =============================================================================
# API Client Fixtures
# =============================================================================


@pytest.fixture
def api_client() -> APIClient:
    """Return unauthenticated API client."""
    return APIClient()


@pytest.fixture
def user(db) -> "User":
    """Create a verified user with default storage quota."""
    user = UserFactory(email_verified=True)
    # Ensure profile has storage quota
    user.profile.storage_quota_bytes = 1024 * 1024 * 1024  # 1GB
    user.profile.total_storage_bytes = 0
    user.profile.save()
    return user


@pytest.fixture
def user_near_quota(db) -> "User":
    """Create a user near their storage quota (only 1MB remaining)."""
    user = UserFactory(email_verified=True)
    user.profile.storage_quota_bytes = 100 * 1024 * 1024  # 100MB
    user.profile.total_storage_bytes = 99 * 1024 * 1024  # 99MB used
    user.profile.save()
    return user


@pytest.fixture
def other_user(db) -> "User":
    """Create another verified user for sharing tests."""
    user = UserFactory(email_verified=True)
    user.profile.storage_quota_bytes = 1024 * 1024 * 1024  # 1GB
    user.profile.total_storage_bytes = 0
    user.profile.save()
    return user


@pytest.fixture
def third_user(db) -> "User":
    """Create a third verified user for multi-user tests."""
    user = UserFactory(email_verified=True)
    user.profile.storage_quota_bytes = 1024 * 1024 * 1024  # 1GB
    user.profile.total_storage_bytes = 0
    user.profile.save()
    return user


@pytest.fixture
def staff_user(db) -> "User":
    """Create a staff user for internal visibility tests."""
    user = UserFactory(email_verified=True, is_staff=True)
    user.profile.storage_quota_bytes = 1024 * 1024 * 1024  # 1GB
    user.profile.total_storage_bytes = 0
    user.profile.save()
    return user


@pytest.fixture
def authenticated_client(user: "User") -> APIClient:
    """Return API client authenticated with JWT token."""
    client = APIClient()
    refresh = RefreshToken.for_user(user)
    client.credentials(HTTP_AUTHORIZATION=f"Bearer {refresh.access_token}")
    return client


@pytest.fixture
def authenticated_client_near_quota(user_near_quota: "User") -> APIClient:
    """Return authenticated client for user near quota."""
    client = APIClient()
    refresh = RefreshToken.for_user(user_near_quota)
    client.credentials(HTTP_AUTHORIZATION=f"Bearer {refresh.access_token}")
    return client


@pytest.fixture
def other_authenticated_client(other_user: "User") -> APIClient:
    """Return API client authenticated as other_user."""
    client = APIClient()
    refresh = RefreshToken.for_user(other_user)
    client.credentials(HTTP_AUTHORIZATION=f"Bearer {refresh.access_token}")
    return client


@pytest.fixture
def staff_authenticated_client(staff_user: "User") -> APIClient:
    """Return API client authenticated as staff user."""
    client = APIClient()
    refresh = RefreshToken.for_user(staff_user)
    client.credentials(HTTP_AUTHORIZATION=f"Bearer {refresh.access_token}")
    return client


# =============================================================================
# Valid File Fixtures
# =============================================================================


@pytest.fixture
def sample_jpeg() -> io.BytesIO:
    """Generate a valid JPEG image file."""
    image = Image.new("RGB", (100, 100), color="red")
    buffer = io.BytesIO()
    image.save(buffer, format="JPEG")
    buffer.seek(0)
    buffer.name = "test_image.jpg"
    return buffer


@pytest.fixture
def sample_jpeg_uploaded(sample_jpeg: io.BytesIO) -> SimpleUploadedFile:
    """Return JPEG as SimpleUploadedFile."""
    return SimpleUploadedFile(
        name="test_image.jpg",
        content=sample_jpeg.read(),
        content_type="image/jpeg",
    )


@pytest.fixture
def sample_png() -> io.BytesIO:
    """Generate a valid PNG image file with transparency."""
    image = Image.new("RGBA", (100, 100), color=(0, 0, 255, 128))
    buffer = io.BytesIO()
    image.save(buffer, format="PNG")
    buffer.seek(0)
    buffer.name = "test_image.png"
    return buffer


@pytest.fixture
def sample_png_uploaded(sample_png: io.BytesIO) -> SimpleUploadedFile:
    """Return PNG as SimpleUploadedFile."""
    return SimpleUploadedFile(
        name="test_image.png",
        content=sample_png.read(),
        content_type="image/png",
    )


@pytest.fixture
def sample_gif() -> io.BytesIO:
    """Generate a valid GIF image file."""
    image = Image.new("P", (50, 50), color=1)
    buffer = io.BytesIO()
    image.save(buffer, format="GIF")
    buffer.seek(0)
    buffer.name = "test_image.gif"
    return buffer


@pytest.fixture
def sample_pdf() -> io.BytesIO:
    """Generate a minimal valid PDF file."""
    # Minimal PDF structure that is recognized by libmagic
    pdf_content = b"""%PDF-1.4
1 0 obj << /Type /Catalog /Pages 2 0 R >> endobj
2 0 obj << /Type /Pages /Kids [] /Count 0 >> endobj
xref
0 3
0000000000 65535 f
0000000009 00000 n
0000000058 00000 n
trailer << /Size 3 /Root 1 0 R >>
startxref
110
%%EOF"""
    buffer = io.BytesIO(pdf_content)
    buffer.name = "test_document.pdf"
    return buffer


@pytest.fixture
def sample_pdf_uploaded(sample_pdf: io.BytesIO) -> SimpleUploadedFile:
    """Return PDF as SimpleUploadedFile."""
    return SimpleUploadedFile(
        name="test_document.pdf",
        content=sample_pdf.read(),
        content_type="application/pdf",
    )


@pytest.fixture
def sample_txt() -> io.BytesIO:
    """Generate a plain text file."""
    content = b"This is a test text file.\nWith multiple lines.\n"
    buffer = io.BytesIO(content)
    buffer.name = "test_document.txt"
    return buffer


@pytest.fixture
def sample_txt_uploaded(sample_txt: io.BytesIO) -> SimpleUploadedFile:
    """Return text file as SimpleUploadedFile."""
    return SimpleUploadedFile(
        name="test_document.txt",
        content=sample_txt.read(),
        content_type="text/plain",
    )


# =============================================================================
# Invalid File Fixtures
# =============================================================================


@pytest.fixture
def empty_file() -> io.BytesIO:
    """Generate an empty file."""
    buffer = io.BytesIO(b"")
    buffer.name = "empty.txt"
    return buffer


@pytest.fixture
def empty_file_uploaded() -> SimpleUploadedFile:
    """Return empty file as SimpleUploadedFile."""
    return SimpleUploadedFile(
        name="empty.txt",
        content=b"",
        content_type="text/plain",
    )


@pytest.fixture
def executable_file() -> io.BytesIO:
    """Generate a file that looks like an ELF executable."""
    # ELF magic bytes followed by padding
    elf_header = b"\x7fELF" + b"\x01\x01\x01\x00" + b"\x00" * 100
    buffer = io.BytesIO(elf_header)
    buffer.name = "malware.jpg"  # Fake .jpg extension
    return buffer


@pytest.fixture
def executable_file_uploaded(executable_file: io.BytesIO) -> SimpleUploadedFile:
    """Return executable as SimpleUploadedFile with fake jpg extension."""
    return SimpleUploadedFile(
        name="totally_not_malware.jpg",
        content=executable_file.read(),
        content_type="image/jpeg",  # Fake content type
    )


@pytest.fixture
def oversized_image() -> io.BytesIO:
    """
    Create a file that exceeds the 25MB image size limit.

    Note: This creates a 26MB file by padding a valid JPEG.
    The file will still be recognized as JPEG by libmagic.
    """
    # Create a small valid JPEG first
    image = Image.new("RGB", (100, 100), color="red")
    buffer = io.BytesIO()
    image.save(buffer, format="JPEG")

    # Get the JPEG data and pad it to exceed 25MB
    jpeg_data = buffer.getvalue()
    # Add padding bytes at the end (after JPEG EOI marker)
    padding_size = (26 * 1024 * 1024) - len(jpeg_data)
    padded_data = jpeg_data + (b"\x00" * padding_size)

    result = io.BytesIO(padded_data)
    result.name = "large_image.jpg"
    return result


@pytest.fixture
def oversized_image_uploaded(oversized_image: io.BytesIO) -> SimpleUploadedFile:
    """Return oversized image as SimpleUploadedFile."""
    content = oversized_image.read()
    return SimpleUploadedFile(
        name="large_image.jpg",
        content=content,
        content_type="image/jpeg",
    )


# =============================================================================
# Extension Mismatch Fixtures
# =============================================================================


@pytest.fixture
def png_with_jpg_extension(sample_png: io.BytesIO) -> io.BytesIO:
    """PNG file with .jpg extension (extension mismatch)."""
    sample_png.seek(0)
    content = sample_png.read()
    buffer = io.BytesIO(content)
    buffer.name = "actually_a_png.jpg"
    return buffer


@pytest.fixture
def pdf_with_txt_extension(sample_pdf: io.BytesIO) -> io.BytesIO:
    """PDF file with .txt extension (extension mismatch)."""
    sample_pdf.seek(0)
    content = sample_pdf.read()
    buffer = io.BytesIO(content)
    buffer.name = "actually_a_pdf.txt"
    return buffer


# =============================================================================
# Malware Scanning Fixtures
# =============================================================================


# EICAR test string - standard antivirus test signature
# This is NOT malware - it's an industry-standard test pattern
EICAR_TEST_STRING = (
    b"X5O!P%@AP[4\\PZX54(P^)7CC)7}$EICAR-STANDARD-ANTIVIRUS-TEST-FILE!$H+H*"
)


@pytest.fixture
def eicar_test_file() -> io.BytesIO:
    """
    Generate EICAR test file for antivirus testing.

    The EICAR test string is a standard pattern recognized
    by all antivirus software as a test signature. It is
    NOT actual malware.
    """
    buffer = io.BytesIO(EICAR_TEST_STRING)
    buffer.name = "eicar_test.txt"
    return buffer


@pytest.fixture
def eicar_test_file_uploaded(eicar_test_file: io.BytesIO) -> SimpleUploadedFile:
    """Return EICAR test file as SimpleUploadedFile."""
    return SimpleUploadedFile(
        name="eicar_test.txt",
        content=eicar_test_file.read(),
        content_type="text/plain",
    )


@pytest.fixture
def mock_clamav_scanner():
    """
    Mock ClamAV scanner for unit tests.

    Provides a mock that simulates ClamAV behavior without
    requiring an actual ClamAV daemon.
    """
    from unittest.mock import MagicMock

    mock = MagicMock()
    mock.ping.return_value = True
    mock.scan_stream.return_value = None  # Clean by default
    mock.scan_file.return_value = None  # Clean by default
    mock.version.return_value = "ClamAV 1.2.0/27234/Mon Dec 11 09:26:33 2024"
    return mock


@pytest.fixture
def media_file_pending_scan(user: "User", sample_jpeg_uploaded: SimpleUploadedFile, db):
    """
    Create a MediaFile awaiting scan.

    The file is in PENDING scan_status, as if just uploaded.
    """
    from pathlib import Path
    from django.conf import settings
    from media.models import MediaFile

    # Create directory for test file
    media_root = Path(settings.MEDIA_ROOT)
    test_dir = media_root / "test_scans"
    test_dir.mkdir(parents=True, exist_ok=True)

    # Write file to disk
    file_path = test_dir / "test_scan_file.jpg"
    file_path.write_bytes(sample_jpeg_uploaded.read())
    sample_jpeg_uploaded.seek(0)

    # Create MediaFile
    media_file = MediaFile.objects.create(
        file="test_scans/test_scan_file.jpg",
        original_filename="test_image.jpg",
        media_type=MediaFile.MediaType.IMAGE,
        mime_type="image/jpeg",
        file_size=sample_jpeg_uploaded.size,
        uploader=user,
        visibility=MediaFile.Visibility.PRIVATE,
        scan_status=MediaFile.ScanStatus.PENDING,
        processing_status=MediaFile.ProcessingStatus.PENDING,
    )

    yield media_file

    # Cleanup
    import shutil

    if file_path.exists():
        file_path.unlink()
    if test_dir.exists():
        shutil.rmtree(test_dir, ignore_errors=True)


# =============================================================================
# Video File Fixtures
# =============================================================================


@pytest.fixture
def sample_mp4() -> io.BytesIO:
    """
    Generate a minimal valid MP4 video file.

    This creates a tiny valid MP4 container using raw bytes.
    It's enough for format detection but not playable.
    For actual video processing tests, use mock_ffmpeg fixture.
    """
    # Minimal ftyp + moov structure for a valid MP4
    # This is the bare minimum to be recognized as MP4
    ftyp = (
        b"\x00\x00\x00\x14"  # box size (20 bytes)
        b"ftyp"  # box type
        b"isom"  # major brand
        b"\x00\x00\x02\x00"  # minor version
        b"isom"  # compatible brand
    )

    # Minimal moov box (empty but valid structure)
    moov = (
        b"\x00\x00\x00\x08"  # box size (8 bytes)
        b"moov"  # box type
    )

    buffer = io.BytesIO(ftyp + moov)
    buffer.name = "test_video.mp4"
    return buffer


@pytest.fixture
def sample_mp4_uploaded(sample_mp4: io.BytesIO) -> SimpleUploadedFile:
    """Return MP4 as SimpleUploadedFile."""
    return SimpleUploadedFile(
        name="test_video.mp4",
        content=sample_mp4.read(),
        content_type="video/mp4",
    )


# =============================================================================
# Document File Fixtures
# =============================================================================


@pytest.fixture
def sample_pdf_with_text() -> io.BytesIO:
    """
    Generate a PDF with extractable text content.

    This creates a minimal PDF that has text that can be extracted
    by pdfplumber for text extraction tests.
    """
    # A simple PDF with actual text content
    pdf_content = b"""%PDF-1.4
1 0 obj << /Type /Catalog /Pages 2 0 R >> endobj
2 0 obj << /Type /Pages /Kids [3 0 R] /Count 1 >> endobj
3 0 obj << /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792]
/Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >> endobj
4 0 obj << /Length 44 >>
stream
BT /F1 12 Tf 100 700 Td (Hello World) Tj ET
endstream
endobj
5 0 obj << /Type /Font /Subtype /Type1 /BaseFont /Helvetica >> endobj
xref
0 6
0000000000 65535 f
0000000009 00000 n
0000000058 00000 n
0000000115 00000 n
0000000266 00000 n
0000000359 00000 n
trailer << /Size 6 /Root 1 0 R >>
startxref
434
%%EOF"""
    buffer = io.BytesIO(pdf_content)
    buffer.name = "test_document_with_text.pdf"
    return buffer


@pytest.fixture
def sample_pdf_with_text_uploaded(
    sample_pdf_with_text: io.BytesIO,
) -> SimpleUploadedFile:
    """Return PDF with text as SimpleUploadedFile."""
    return SimpleUploadedFile(
        name="test_document_with_text.pdf",
        content=sample_pdf_with_text.read(),
        content_type="application/pdf",
    )


@pytest.fixture
def sample_docx() -> io.BytesIO:
    """
    Generate a minimal Word document (DOCX).

    Uses python-docx if available, otherwise creates a minimal valid DOCX.
    """
    try:
        from docx import Document

        doc = Document()
        doc.add_paragraph("Test document content")
        doc.add_paragraph("With multiple paragraphs for testing.")

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        buffer.name = "test_document.docx"
        return buffer
    except ImportError:
        # Fallback: Create minimal DOCX structure (ZIP with XML)
        import zipfile

        buffer = io.BytesIO()
        with zipfile.ZipFile(buffer, "w", zipfile.ZIP_DEFLATED) as zf:
            # Minimal [Content_Types].xml
            content_types = b"""<?xml version="1.0" encoding="UTF-8"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
<Default Extension="xml" ContentType="application/xml"/>
<Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>
</Types>"""
            zf.writestr("[Content_Types].xml", content_types)

            # Minimal _rels/.rels
            rels = b"""<?xml version="1.0" encoding="UTF-8"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>
</Relationships>"""
            zf.writestr("_rels/.rels", rels)

            # Minimal word/document.xml
            document = b"""<?xml version="1.0" encoding="UTF-8"?>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
<w:body><w:p><w:r><w:t>Test content</w:t></w:r></w:p></w:body>
</w:document>"""
            zf.writestr("word/document.xml", document)

        buffer.seek(0)
        buffer.name = "test_document.docx"
        return buffer


@pytest.fixture
def sample_docx_uploaded(sample_docx: io.BytesIO) -> SimpleUploadedFile:
    """Return DOCX as SimpleUploadedFile."""
    return SimpleUploadedFile(
        name="test_document.docx",
        content=sample_docx.read(),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


@pytest.fixture
def sample_xlsx() -> io.BytesIO:
    """
    Generate a minimal Excel spreadsheet (XLSX).

    Uses openpyxl if available, otherwise creates a minimal valid XLSX.
    """
    try:
        from openpyxl import Workbook

        wb = Workbook()
        ws = wb.active
        ws.title = "Test Sheet"
        ws["A1"] = "Test"
        ws["B1"] = "Data"
        ws["A2"] = 123
        ws["B2"] = 456

        buffer = io.BytesIO()
        wb.save(buffer)
        buffer.seek(0)
        buffer.name = "test_spreadsheet.xlsx"
        return buffer
    except ImportError:
        # Fallback: Create minimal XLSX structure (ZIP with XML)
        import zipfile

        buffer = io.BytesIO()
        with zipfile.ZipFile(buffer, "w", zipfile.ZIP_DEFLATED) as zf:
            # Minimal [Content_Types].xml
            content_types = b"""<?xml version="1.0" encoding="UTF-8"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
<Default Extension="xml" ContentType="application/xml"/>
<Override PartName="/xl/workbook.xml" ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet.main+xml"/>
</Types>"""
            zf.writestr("[Content_Types].xml", content_types)

            # Minimal _rels/.rels
            rels = b"""<?xml version="1.0" encoding="UTF-8"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="xl/workbook.xml"/>
</Relationships>"""
            zf.writestr("_rels/.rels", rels)

            # Minimal xl/workbook.xml
            workbook = b"""<?xml version="1.0" encoding="UTF-8"?>
<workbook xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main">
<sheets><sheet name="Sheet1" sheetId="1" r:id="rId1" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"/></sheets>
</workbook>"""
            zf.writestr("xl/workbook.xml", workbook)

        buffer.seek(0)
        buffer.name = "test_spreadsheet.xlsx"
        return buffer


@pytest.fixture
def sample_xlsx_uploaded(sample_xlsx: io.BytesIO) -> SimpleUploadedFile:
    """Return XLSX as SimpleUploadedFile."""
    return SimpleUploadedFile(
        name="test_spreadsheet.xlsx",
        content=sample_xlsx.read(),
        content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )


# =============================================================================
# Processing Test Fixtures
# =============================================================================


@pytest.fixture
def media_file_for_processing(
    user: "User", sample_jpeg_uploaded: SimpleUploadedFile, db
):
    """
    Create a MediaFile ready for processing (scan complete).

    The file is in CLEAN scan_status and PENDING processing_status.
    """
    from pathlib import Path
    from django.conf import settings
    from media.models import MediaFile

    # Create directory for test file
    media_root = Path(settings.MEDIA_ROOT)
    test_dir = media_root / "test_processing"
    test_dir.mkdir(parents=True, exist_ok=True)

    # Write file to disk
    file_path = test_dir / "test_process_file.jpg"
    file_path.write_bytes(sample_jpeg_uploaded.read())
    sample_jpeg_uploaded.seek(0)

    # Create MediaFile
    media_file = MediaFile.objects.create(
        file="test_processing/test_process_file.jpg",
        original_filename="test_image.jpg",
        media_type=MediaFile.MediaType.IMAGE,
        mime_type="image/jpeg",
        file_size=sample_jpeg_uploaded.size,
        uploader=user,
        visibility=MediaFile.Visibility.PRIVATE,
        scan_status=MediaFile.ScanStatus.CLEAN,
        processing_status=MediaFile.ProcessingStatus.PENDING,
    )

    yield media_file

    # Cleanup
    import shutil

    if file_path.exists():
        file_path.unlink()
    if test_dir.exists():
        shutil.rmtree(test_dir, ignore_errors=True)


@pytest.fixture
def mock_ffmpeg():
    """
    Mock FFmpeg/FFprobe for video processing tests.

    Simulates FFmpeg behavior without requiring actual video files
    or the FFmpeg binary.
    """
    from unittest.mock import MagicMock, patch

    # Mock ffprobe output for a typical video
    ffprobe_output = {
        "streams": [
            {
                "codec_type": "video",
                "codec_name": "h264",
                "width": 1920,
                "height": 1080,
                "r_frame_rate": "30000/1001",
                "duration": "60.0",
            },
            {
                "codec_type": "audio",
                "codec_name": "aac",
            },
        ],
        "format": {
            "duration": "60.0",
            "bit_rate": "5000000",
        },
    }

    mock_run = MagicMock()
    mock_run.return_value.returncode = 0
    mock_run.return_value.stdout = __import__("json").dumps(ffprobe_output)
    mock_run.return_value.stderr = ""

    with patch("subprocess.run", mock_run):
        yield mock_run


@pytest.fixture
def mock_libreoffice():
    """
    Mock LibreOffice for document conversion tests.

    Simulates LibreOffice headless conversion without requiring
    the actual LibreOffice installation.
    """
    from unittest.mock import MagicMock, patch

    mock_run = MagicMock()
    mock_run.return_value.returncode = 0
    mock_run.return_value.stdout = ""
    mock_run.return_value.stderr = ""

    with patch("subprocess.run", mock_run):
        yield mock_run
